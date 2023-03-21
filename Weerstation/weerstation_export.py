from docx import Document
from datetime import datetime

"""Define a function that can check whether something is the same as the request"""
"""If this is the case, append it"""


def check_append(input, check, append):
    if input.text == check.strip():
        input.text = append


"""Create a function that creates a Word file based on the variables within the loc.data"""


def create_word_export(loc, location):
    time = datetime.now()
    time = time.strftime('%d-%m-%Y %Hh%Mm')
    # Set up filename and the name of the template doc
    filename = location + ' ' + time + '.docx'
    doc = Document('Weather_template.docx')
    # Make a temporary lib where any data can be put into
    data = {}

    # reformat everything into '_{}_' and #_{}_
    for items in loc.data:
        new = items.replace(' ', '_').lower()
        value = f'#_{new}_'
        data[value] = str(loc.data[items])

    # Check for any key term in the paragraphs in the document
    for items in doc.paragraphs:
        for var in data:
            check_append(items, var, data[var])

    # check if any tables are in the document and if that's the case, check the text within the cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for var in data:
                    check_append(cell, var, data[var])

    # Save the file in the 'word_files' dic and as the filename mentioned above
    doc.save('word_files/' + filename)
